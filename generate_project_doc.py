#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Export project tree + source files into a single Word (.docx).

Usage:
  python export_project_to_word.py --root . --out project_dump.docx
  python export_project_to_word.py --root . --out dump.docx --ext .py .yaml .yml .json .sh
"""

import os
import sys
import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


DEFAULT_EXCLUDE_DIRS = {
    ".git", ".hg", ".svn",
    "__pycache__", ".mypy_cache", ".pytest_cache",
    ".idea", ".vscode",
    "venv", ".venv", "env", ".env",
    "build", "dist", ".eggs",
    "logs", "log", "outputs", "output", "results", "runs",
    "data", "datasets", "assets", "tmp", "cache",
    "checkpoints", "weights",
}

DEFAULT_EXCLUDE_FILE_EXTS = {
    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp",
    ".mp4", ".mov", ".avi", ".mkv",
    ".pt", ".pth", ".ckpt", ".onnx", ".tflite",
    ".npz", ".npy",
    ".zip", ".tar", ".gz", ".7z", ".rar",
    ".so", ".dll", ".dylib",
    ".pdf",
}

DEFAULT_INCLUDE_FILES = {
    "pyproject.toml", "requirements.txt", "environment.yml",
    "setup.py", "setup.cfg", "Pipfile", "Pipfile.lock",
    "README.md", "README.rst",
}


def set_monospace(run, font_name="Consolas", font_size=9):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)  # for CJK
    run.font.size = Pt(font_size)


def add_code_block(doc: Document, text: str):
    """
    Insert a preformatted-like block: single paragraph, monospace, line breaks preserved.
    """
    p = doc.add_paragraph()
    r = p.add_run()
    set_monospace(r, font_name="Consolas", font_size=9)
    # Preserve newlines
    lines = text.splitlines()
    for i, line in enumerate(lines):
        if i > 0:
            r.add_break()  # line break
        r.add_text(line)


def build_tree(root: Path, exclude_dirs: set, max_entries: int = 20000) -> str:
    """
    Build a 'tree'-like text without calling external commands.
    """
    root = root.resolve()
    lines = [str(root)]
    entries_count = 0

    def is_excluded_dir(p: Path) -> bool:
        return p.name in exclude_dirs

    def walk_dir(current: Path, prefix: str = ""):
        nonlocal entries_count
        if entries_count >= max_entries:
            lines.append(prefix + "└── [TRUNCATED: too many entries]")
            return

        try:
            children = sorted(list(current.iterdir()), key=lambda x: (x.is_file(), x.name.lower()))
        except PermissionError:
            lines.append(prefix + "└── [PermissionError]")
            return

        # filter excluded dirs
        filtered = []
        for c in children:
            if c.is_dir() and is_excluded_dir(c):
                continue
            filtered.append(c)

        for idx, c in enumerate(filtered):
            if entries_count >= max_entries:
                lines.append(prefix + "└── [TRUNCATED: too many entries]")
                return
            entries_count += 1

            is_last = (idx == len(filtered) - 1)
            connector = "└── " if is_last else "├── "
            lines.append(prefix + connector + c.name)

            if c.is_dir():
                extension = "    " if is_last else "│   "
                walk_dir(c, prefix + extension)

    walk_dir(root, "")
    return "\n".join(lines)


def should_include_file(path: Path, exts: set, exclude_file_exts: set) -> bool:
    if path.name in DEFAULT_INCLUDE_FILES:
        return True
    if path.suffix.lower() in exclude_file_exts:
        return False
    return path.suffix.lower() in exts


def iter_files(root: Path, exclude_dirs: set, exts: set, exclude_file_exts: set):
    root = root.resolve()
    for dirpath, dirnames, filenames in os.walk(root):
        # prune dirs in-place
        dirnames[:] = [d for d in dirnames if d not in exclude_dirs]
        for fn in sorted(filenames):
            p = Path(dirpath) / fn
            if should_include_file(p, exts, exclude_file_exts):
                yield p


def read_text_file(path: Path, max_bytes: int, max_lines: int):
    data = path.read_bytes()
    truncated = False

    if len(data) > max_bytes:
        data = data[:max_bytes]
        truncated = True

    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("utf-8", errors="replace")

    lines = text.splitlines()
    if len(lines) > max_lines:
        lines = lines[:max_lines]
        truncated = True

    return "\n".join(lines), truncated


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--root", type=str, default=".", help="project root")
    ap.add_argument("--out", type=str, default="project_dump.docx", help="output docx path")
    ap.add_argument("--ext", nargs="*", default=[".py"], help="extensions to include, e.g. .py .yaml .yml")
    ap.add_argument("--exclude-dirs", nargs="*", default=sorted(DEFAULT_EXCLUDE_DIRS), help="dir names to exclude")
    ap.add_argument("--max-bytes", type=int, default=300_000, help="max bytes per file before truncation")
    ap.add_argument("--max-lines", type=int, default=5000, help="max lines per file before truncation")
    ap.add_argument("--max-tree-entries", type=int, default=20000, help="cap project tree entries")
    args = ap.parse_args()

    root = Path(args.root).resolve()
    out_path = Path(args.out).resolve()

    exts = {e if e.startswith(".") else ("." + e) for e in args.ext}
    exclude_dirs = set(args.exclude_dirs)

    doc = Document()

    # Title
    doc.add_heading(f"Project Dump: {root.name}", level=0)
    doc.add_paragraph(f"Root: {root}")
    doc.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Included extensions: {', '.join(sorted(exts))}")
    doc.add_paragraph(f"Excluded dirs: {', '.join(sorted(exclude_dirs))}")
    doc.add_paragraph(f"Per-file truncation: max_bytes={args.max_bytes}, max_lines={args.max_lines}")

    # Tree
    doc.add_heading("Project Tree", level=1)
    tree_text = build_tree(root, exclude_dirs=exclude_dirs, max_entries=args.max_tree_entries)
    add_code_block(doc, tree_text)

    # Files
    doc.add_heading("Files", level=1)
    file_count = 0
    total_bytes = 0

    for p in iter_files(root, exclude_dirs=exclude_dirs, exts=exts, exclude_file_exts=DEFAULT_EXCLUDE_FILE_EXTS):
        rel = p.relative_to(root)
        try:
            size = p.stat().st_size
        except OSError:
            continue

        file_count += 1
        total_bytes += size

        doc.add_heading(str(rel), level=2)
        doc.add_paragraph(f"Size: {size} bytes")

        try:
            content, truncated = read_text_file(p, max_bytes=args.max_bytes, max_lines=args.max_lines)
        except Exception as e:
            doc.add_paragraph(f"[READ ERROR] {type(e).__name__}: {e}")
            continue

        if truncated:
            doc.add_paragraph("[TRUNCATED] content was cut by size/line limits.")

        add_code_block(doc, content)

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Included files: {file_count}")
    doc.add_paragraph(f"Total raw size (before truncation): {total_bytes} bytes")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Wrote: {out_path}")


if __name__ == "__main__":
    main()
